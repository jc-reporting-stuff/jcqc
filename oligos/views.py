from django.forms.models import formset_factory
from django.http.response import HttpResponseRedirect
from django.shortcuts import redirect, render
from django.views.generic import ListView, CreateView, FormView, TemplateView, DetailView
from django.urls import reverse_lazy, reverse
from django.http import HttpResponse, FileResponse
from django.db.models import Max
from django.contrib import messages
from django.utils.decorators import method_decorator
from django.views.generic.base import View
from django.utils.timezone import localtime, now, make_aware
from django.core.paginator import Paginator
from django.conf import settings

from oligos.forms import EasyOrderForm, EnterODForm, OligoInitialForm, OligoOrderForm, IdRangeForm, DateRangeForm, OligoTextSearch, PriceForm
from .models import Oligo, OliPrice
from accounts.models import Account, User
from core.decorators import owner_or_staff, user_has_accounts

from docx import Document

import io
import pytz
import datetime
import re


class ClientOrderListView(ListView):
    template_name = 'oligos/client_view_all.html'
    context_object_name = 'orders'

    def get_queryset(self):
        qs = Oligo.objects.filter(submitter=self.request.user) | Oligo.objects.filter(
            account__owner__username=self.request.user)
        return qs


class OligoDetailView(DetailView):
    model = Oligo

    context_object_name = 'oligo'

    def get_queryset(self):
        user = self.request.user
        qs = super().get_queryset()
        return qs.filter(submitter__id=user.id) | qs.filter(account__owner__id=user.id)


@method_decorator(owner_or_staff, name='dispatch')
class UserListView(ListView):
    context_object_name = 'oligos'
    paginate_by = settings.PAGINATE_HISTORY_LENGTH
    template_name = 'oligos/user_history.html'

    def get_queryset(self):
        username = self.kwargs.get('username')
        qs = Oligo.objects.all()
        qs = qs.filter(submitter__username=username) | qs.filter(
            account__owner__username=username)
        return qs


@method_decorator(user_has_accounts, name='dispatch')
class OligoNewTypeView(FormView):
    form_class = OligoInitialForm
    template_name = 'oligos/order_method.html'
    success_url = reverse_lazy('oligos:order_create')

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()

        user = self.request.user
        approved_accounts = user.get_financial_accounts()
        approved_choices = [((account.id), (': '.join(
            [account.code, account.comment]))) for account in approved_accounts]
        kwargs['account_choices'] = approved_choices
        kwargs['initial_oligos'] = 1
        return kwargs


class OligoCreateView(CreateView):
    template_name = 'oligos/order_new.html'
    model = Oligo
    fields = ['name', 'sequence', 'order_id']

    def get(self, request, *args, **kwargs):
        # If coming from the "how many oligos page", this will be carried out.
        if request.GET.get('account_id'):
            account_id = request.GET.get('account_id')
            oligo_count = request.GET.get('oligo_count')

            NewOligoFormset = formset_factory(
                OligoOrderForm, extra=int(oligo_count))
            formset = NewOligoFormset()

            response = render(request, 'oligos/order_new.html',
                              context={'formset': formset, 'account_id': account_id})
            return HttpResponse(response)
        return HttpResponseRedirect(reverse_lazy('oligos:order_quantity'))

    def post(self, request, *args, **kwargs):
        OligoFormset = formset_factory(OligoOrderForm)
        formset = OligoFormset(request.POST)

        # Make sure that all required fields are filled out.
        num_forms = 0
        for form in formset:
            if form.has_changed():
                num_forms += 1
            if not form.is_valid():
                return HttpResponse(render(request, 'oligos/order_new.html', context={'formset': formset}))

        # If nothing was changed in any form, send user back to choose number to order.
        if num_forms == 0:
            return HttpResponseRedirect(reverse_lazy('oligos:order_quantity'))

        # If the 'preview' button has been pressed, re-open the template in preview mode
        previewing = request.POST.get('previewing')

        account_id = request.POST.get('account_id')

        if previewing:
            account = Account.objects.get(id=account_id)
            sequences = []

            for form in formset:
                spaced_sequence = ' '.join(form['sequence'].value()[i:i+3]
                                           for i in range(0, len(form['sequence'].value()), 3))
                sequences.append(spaced_sequence)

            context = {
                'formset': formset,
                'previewing': True,
                'account_id': account_id,
                'sequences': sequences,
                'account': account
            }
            response = render(self.request, 'oligos/order_new.html', context)
            return HttpResponse(response)

        return self.form_valid(formset)

    def form_valid(self, formset):
        cd = []
        for form in formset:
            cleaned_data = form.clean()
            if cleaned_data:
                cd.append(cleaned_data)

        max_order = Oligo.objects.aggregate(Max('order_id'))
        if not max_order['order_id__max']:
            max_order_id = 0
        else:
            max_order_id = max_order['order_id__max']
        order_number = max_order_id + 1

        submitter = User.objects.get(id=self.request.user.id)
        account = Account.objects.get(id=self.request.POST.get('account_id'))

        for oligo in cd:
            Oligo.objects.create(
                name=oligo['name'], sequence=oligo['sequence'].upper(), order_id=order_number,
                account=account, submitter=submitter, modification=oligo['modification'],
                scale=oligo['scale'], purity=oligo['purity']
            )

        request = self.request
        messages.success(request, r'Ordered Successfully.')

        return HttpResponseRedirect(reverse_lazy('oligos:client_order_list'))


class OligoEasyOrder(FormView):
    form_class = EasyOrderForm
    template_name = 'oligos/easy_order.html'
    success_url = reverse_lazy('oligos:easy_order')
    context_object_name = 'context'

    def get(self, request):
        # If coming from the "how many oligos page", this will be carried out.
        if request.GET.get('account_id'):
            account_id = request.GET.get('account_id')

            response = render(request, 'oligos/easy_order.html',
                              context={'form': EasyOrderForm, 'account_id': account_id})
            return HttpResponse(response)
        return HttpResponseRedirect(reverse_lazy('oligos:order_quantity'))

    def form_valid(self, form):
        cd = form.cleaned_data
        account_id = self.request.POST.get('account_id') if self.request.POST.get(
            'account_id') else self.request.GET.get('account_id')

        # Process the oligos from the textarea field.
        # Split them by line, then use regular expressions to split into name, sequence groups.
        oligo_text = form['oligos'].value()
        oligos_text_split = oligo_text.split('\n')
        ex = r'(.+?)[\t;,]\s*([ACGTRYMWSKDHBVNacgtrymwskdhbvn\s]+)'
        oligo_groups = []
        for line in oligos_text_split:
            if line and line != '\r':
                groups = re.match(ex, line).groups()
                sequence = groups[1].upper().strip().replace(' ', '')
                display_sequence = ' '.join(sequence[i:i+3]
                                            for i in range(0, len(sequence), 3))
                current_oligo = {
                    'name': groups[0], 'sequence': sequence, 'display_sequence': display_sequence}
                oligo_groups.append(current_oligo)

        account = Account.objects.get(id=account_id)

        context = {
            'account': account,
            'data': cd,
            'oligos': oligo_groups,
            'previewing': True
        }

        # Pass everything along to the preview view.
        response = render(
            self.request, 'oligos/easy_preview.html', context=context)
        return HttpResponse(response)


class OligoEasySubmitView(TemplateView):
    template_name = 'oligos/easy_preview.html'

    def post(self, request):
        data = request.POST
        submitter = User.objects.get(id=self.request.user.id)
        account = Account.objects.get(id=data['account_id'])

        # Build a list of oligos from the form data.
        oligos = []
        for i in range(int(data['oligo_count'])):
            object = {}
            form_oligo_name = 'oligo-' + str(i) + '-name'
            form_sequence_name = 'oligo-' + str(i) + '-sequence'
            object['name'] = data[form_oligo_name]
            object['sequence'] = data[form_sequence_name]
            oligos.append(object)

        # Cycle through list of oligos, add necessary fields and save.
        max_order = Oligo.objects.aggregate(Max('order_id'))
        order_number = max_order['order_id__max'] + 1

        for oligo in oligos:
            try:
                object = Oligo.objects.create(
                    name=oligo['name'], sequence=oligo['sequence'], order_id=order_number,
                    account=account, submitter=submitter, modification=data['modification'],
                    scale=data['scale'], purity=data['purity']
                )
            except:
                messages.add_message(
                    request, r'Saving oligo {oligo.name} failed! Check your values and try again.')

        return HttpResponseRedirect(reverse_lazy('oligos:client_order_list'))


class OligoAdminHomeView(View):
    def get(self, request):
        return render(request, 'oligos/admin_home.html')


class OligoTodayListView(ListView):
    model = Oligo
    context_object_name = 'oligos'
    template_name = 'oligos/todays_oligos.html'

    def get_queryset(self, *args, **kwargs):
        today = localtime(now()).date()
        return Oligo.objects.filter(created_at__gte=today).order_by('id')

    def get_context_data(self, **kwargs):
        context = super(OligoTodayListView, self).get_context_data(**kwargs)
        context['title'] = "Today's Oligos"
        return context


def get_redirect_url(search, low=None, high=None, text='', client=''):
    base_url = reverse_lazy('oligos:search_results')
    low_str = f'&low={low}' if low else ''
    high_str = f'&high={high}' if high else ''
    text_str = f'&text={text}' if text else ''
    client_str = f'&client={client}' if client else ''
    return f'{base_url}?query={search}{low_str}{high_str}{text_str}{client_str}'


def get_search_queries(search_type):
    if search_type == 'range':
        return ['oligo', 'date', 'order']
    elif search_type == 'text':
        return ['name', 'sequence']
    else:
        return []


class OligoSearchView(View):
    def get(self, request):
        if request.GET.get('search'):
            args = request.GET
            search = args.get('search')
            range_searches = get_search_queries('range')
            text_searches = get_search_queries('text')
            if search in range_searches:
                low = args.get(f'{search}-low')
                high = args.get(f'{search}-high')
                client = args.get(f'{search}-client')
                return redirect(get_redirect_url(search=search, low=low, high=high, client=client))
            elif search in text_searches:
                text = args.get(f'{search}-text')
                return redirect(get_redirect_url(search=search, text=text))

        max_oligo = Oligo.objects.all().order_by('-id').first()
        id_range_form = IdRangeForm(
            prefix="oligo", initial={'high': max_oligo.id})
        date_range_form = DateRangeForm(prefix="date")
        order_range_form = IdRangeForm(
            prefix="order", initial={'high': max_oligo.order_id})
        oligo_name_form = OligoTextSearch(prefix="name")
        oligo_sequence_form = OligoTextSearch(prefix="sequence")

        context = {
            'id_range_form': id_range_form,
            'date_range_form': date_range_form,
            'order_range_form': order_range_form,
            'oligo_name_form': oligo_name_form,
            'oligo_sequence_form': oligo_sequence_form
        }
        return render(request, 'oligos/searches.html', context=context)


def create_aware_date(form_date_string):
    time_zone = pytz.timezone('Canada/Eastern')
    split_date = form_date_string.split('-')
    naive_date = datetime.datetime(
        int(split_date[0]), int(split_date[1]), int(split_date[2]))
    return make_aware(naive_date, time_zone)


class OligoListView(View):
    def get(self, request):
        args = request.GET

        ### Error-checking -- make sure all necessary search params exist and are valid ###
        params = ['query', 'low', 'high', 'text', 'client', 'page']
        data = {}
        for param in params:
            data[param] = args.get(param) if args.get(param) else None
        if not (data['query'] and ((data['low'] and data['high']) or data['text'])) and not data['page']:
            messages.warning(
                request, 'Missing search parameter(s), try again please.')
            return redirect(reverse('oligos:search'))

        if data['low'] and data['high']:
            if data['low'] > data['high']:
                temp = data['low']
                data['low'] = data['high']
                data['high'] = temp
                messages.info(
                    request, 'Low value greater than high value: Values were swapped.')

        queryset = None

        if data['query'] == 'oligo':
            queryset = Oligo.objects.filter(
                id__gte=data['low'], id__lte=data['high'])

        elif data['query'] == 'date':
            start_date = create_aware_date(data['low'])
            end_date = create_aware_date(
                data['high']) + datetime.timedelta(days=1)

            client = data['client'] or False
            queryset = Oligo.objects.filter(
                created_at__gte=start_date, created_at__lte=end_date)
            if client != '0':
                queryset = queryset.filter(submitter__id=int(data['client']))

        elif data['query'] == 'order':
            queryset = Oligo.objects.filter(
                order_id__gte=data['low'], order_id__lte=data['high'])

        elif data['query'] == 'name':
            queryset = Oligo.objects.filter(name__icontains=data['text'])

        elif data['query'] == 'sequence':
            stripped_sequence = data['text'].strip().replace(' ', '')
            queryset = Oligo.objects.filter(
                sequence__icontains=stripped_sequence)

        queryset = queryset.order_by("id")

        paginator = Paginator(queryset, settings.PAGINATE_SEARCH_LENGTH)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        current_path = request.get_full_path()

        context = {
            'page_obj': page_obj,
            'data': data,
            'current_path': current_path,
        }

        return render(request, 'oligos/admin_list.html', context=context)


def get_checked_oligos_from_data(list_of_dicts):
    oligos_to_update = []
    for key in list_of_dicts.keys():
        try:
            checkbox = re.match('^(\d+)-checkbox$', key).groups()[0]
            oligos_to_update.append(checkbox)
        except:
            pass
    return oligos_to_update


class OligoListActionsView(View):
    def post(self, request):
        action = request.POST.get('button')

        post_data = dict(request.POST.lists())

        if action == 'update-delivery':
            oligos_to_update = get_checked_oligos_from_data(post_data)
            for oligo_to_update in oligos_to_update:
                if request.POST.get(f'{oligo_to_update}-delivery-date'):
                    date = request.POST.get(f'{oligo_to_update}-delivery-date')
                    aware_date = create_aware_date(date)
                    oligo = Oligo.objects.get(id=oligo_to_update)
                    oligo.delivery_date = aware_date
                    oligo.save()
                else:
                    oligo = Oligo.objects.get(id=oligo_to_update)
                    oligo.delivery_date = None
                    oligo.save()
            messages.success(request, f'Delivery dates saved.')

        elif action == 'update-all-delivery':
            delivery_date = request.POST.get('all-delivery-date')
            if delivery_date:
                aware_delivery_date = create_aware_date(delivery_date)
            else:
                aware_delivery_date = None
            for key in post_data.keys():
                try:
                    oligo_id = re.match(
                        '^(\d+)-delivery-date$', key).groups()[0]
                except:
                    continue
                oligo = Oligo.objects.get(id=oligo_id)
                oligo.delivery_date = aware_delivery_date
                oligo.save()
            messages.success(request, f'Delivery dates saved successfully.')

        elif action == 'make-report':
            oligos_to_update = get_checked_oligos_from_data(post_data)
            oligos = []
            for oligo in oligos_to_update:
                oligo_object = Oligo.objects.get(id=oligo)
                oligos.append(oligo_object)

            for oligo in oligos:
                if oligo.OD_reading:
                    oligo.micrograms95 = round(
                        oligo.micrograms_per_microliter * 95, 1)
                    oligo.nmols95 = round(
                        oligo.pmol_per_microliter * 95 / 1000, 1)
                    oligo.micrograms195 = round(
                        oligo.micrograms_per_microliter * 195, 1)
                    oligo.nmols195 = round(
                        oligo.pmol_per_microliter * 195 / 1000, 1)
            return render(request, 'oligos/report.html', context={'oligos': oligos})

        elif action == 'create-labels':
            oligos_to_label = get_checked_oligos_from_data(post_data)
            if len(oligos_to_label) > 80:
                messages.warning(
                    request, f'Can only make up to 80 labels at a time. Labels made using first 80 by ID ascending.')
                oligos_to_label = oligos_to_label[:80]
            return CreateLabels(oligos_to_label)

        redirect_url = request.POST.get('return-url')
        return HttpResponseRedirect(redirect_url)


def ReportOrderView(request):
    if not request.GET.get('order_id'):
        return HttpResponseRedirect(reverse('oligos:search'))
    order_id = request.GET.get('order_id')
    oligos = list(Oligo.objects.filter(order_id=order_id))
    for oligo in oligos:
        oligo.micrograms95 = round(oligo.micrograms_per_microliter * 95, 1)
        oligo.nmols95 = round(oligo.pmol_per_microliter * 95 / 1000, 1)
        oligo.micrograms195 = round(oligo.micrograms_per_microliter * 195, 1)
        oligo.nmols195 = round(oligo.pmol_per_microliter * 195 / 1000, 1)
    return render(request, 'oligos/report.html', context={'oligos': oligos})


class BillingView(View):
    def get(self, request):
        max_oligo = Oligo.objects.all().order_by('-id').first()
        range_form = IdRangeForm(initial={'high': max_oligo.id})

        prices = OliPrice.objects.filter(current=True).last()
        price_form = PriceForm(instance=prices, prefix="price")

        context = {
            'range_form': range_form,
            'price_form': price_form
        }
        return render(request, 'oligos/billing.html', context=context)

    def post(self, request):
        range_form = IdRangeForm(request.POST)
        price_form = PriceForm(request.POST, prefix="price")
        if not range_form.is_valid() or not price_form.is_valid():
            context = {
                'range_form': range_form,
                'price_form': price_form
            }
            return render(request, 'oligos/billing.html', context=context)

        price_data = price_form.cleaned_data
        oligo_range = range_form.cleaned_data
        oligos = Oligo.objects.filter(
            id__gte=oligo_range['low']).filter(id__lte=oligo_range['high']).order_by('account__owner__id', 'account__id')

        class ReportObject:
            def __init__(self):
                self.order_dates = []
                self.degen_order_dates = []
                self.submitter_id = ''
                self.account_id = ''
                self.oligo_count = 0
                self.degen_oligo_ids = []
                self.nmol40_count = 0
                self.nmol200_count = 0
                self.nmol1000_count = 0
                self.degen40_count = 0
                self.degen200_count = 0
                self.degen1000_count = 0
                self.desalt_count = 0
                self.cartridge_count = 0

            def add_count(self, scale, degen=False):
                if not degen:
                    if scale == '40':
                        self.nmol40_count += 1
                    elif scale == '200':
                        self.nmol200_count += 1
                    elif scale == '1000':
                        self.nmol1000_count += 1
                else:
                    if scale == '40':
                        self.degen40_count += 1
                    elif scale == '200':
                        self.degen200_count += 1
                    elif scale == '1000':
                        self.degen1000_count += 1

        billing_rows = []
        row_counter = 0
        this_row = None

        ## Build up pricing structure ###
        for oligo in oligos:
            # If the user has changed, we'll add an object to the list and switch to it.
            user_has_changed = False
            account_has_changed = False
            if len(billing_rows) != 0:
                user_has_changed = oligo.submitter.id != this_row.submitter_id
                account_has_changed = oligo.account.id != this_row.account_id

            # If the user has changed, add a new row to the table.
            if len(billing_rows) == 0:
                x = ReportObject()
                billing_rows.append(x)
                this_row = billing_rows[row_counter]
                this_row.submitter_id = oligo.submitter.id
                this_row.account_id = oligo.account.id

            elif user_has_changed or account_has_changed:
                x = ReportObject()
                billing_rows.append(x)
                row_counter += 1
                this_row = billing_rows[row_counter]
                this_row.submitter_id = oligo.submitter.id
                this_row.account_id = oligo.account.id

            ## Start processing of actual oligo ###
            this_row.oligo_count += 1

            non_degen_bases = ['A', 'C', 'G', 'T']
            oligo_scale = oligo.scale.split(' ')[0]
            if oligo_scale == '1':
                oligo_scale = '1000'
            for base in oligo.sequence:
                if base in non_degen_bases:
                    this_row.add_count(oligo_scale)
                else:
                    this_row.add_count(oligo_scale, degen=True)
                    if oligo.id not in this_row.degen_oligo_ids:
                        this_row.degen_oligo_ids.append(oligo.id)
                    if oligo.created_at.date() not in this_row.degen_order_dates:
                        this_row.degen_order_dates.append(
                            oligo.created_at.date())

            if oligo.purity == 'desalted':
                this_row.desalt_count += 1
            elif oligo.purity == 'cartridge':
                this_row.cartridge_count += 1

            if oligo.created_at.date() not in this_row.order_dates:
                this_row.order_dates.append(oligo.created_at.date())

        class RowObject:
            def __init__(self, submitter, account):
                self.submitter = submitter
                self.account = account

        billing_context = []
        for row in billing_rows:
            submitter = User.objects.get(id=row.submitter_id)
            account = Account.objects.get(id=row.account_id)
            row_object = RowObject(submitter, account)
            row_object.order_date = row.order_dates
            row_object.supervisor = account.owner.display_name
            row_object.client_name = submitter.display_name
            row_object.department = submitter.department
            row_object.account = account.code
            row_object.oligo_count = row.oligo_count
            row_object.nmol40 = row.nmol40_count
            row_object.price40 = price_data['scale_40_base']
            row_object.nmol200 = row.nmol200_count
            row_object.price200 = price_data['scale_200_base']
            row_object.nmol1000 = row.nmol1000_count
            row_object.price1000 = price_data['scale_1000_base']

            row_object.degen_order_date = row.degen_order_dates
            row_object.degen_oligo_count = len(row.degen_oligo_ids)
            row_object.nmol40_degen = row.degen40_count
            row_object.price40_degen = price_data['degenerate_40_base']
            row_object.nmol200_degen = row.degen200_count
            row_object.price200_degen = price_data['degenerate_200_base']
            row_object.nmol1000_degen = row.degen1000_count
            row_object.price1000_degen = price_data['degenerate_1000_base']

            row_object.desalt_count = row.desalt_count
            row_object.desalt_price = price_data['desalt_fee']
            row_object.cartridge_count = row.cartridge_count
            row_object.cartridge_price = price_data['cartridge_fee']
            row_object.setup_count = len(row.order_dates)
            row_object.setup_price = price_data['setup_fee']

            price40 = row.nmol40_count * price_data['scale_40_base']
            price200 = row.nmol200_count * price_data['scale_200_base']
            price1000 = row.nmol1000_count * price_data['scale_1000_base']

            price_desalt = row.desalt_count * price_data['desalt_fee']
            price_cartridge = row.cartridge_count * price_data['cartridge_fee']
            price_setup = len(row.order_dates) * price_data['setup_fee']

            price40_degen = row.degen40_count * \
                price_data['degenerate_40_base']
            price200_degen = row.degen200_count * \
                price_data['degenerate_200_base']
            price1000_degen = row.degen1000_count * \
                price_data['degenerate_1000_base']

            row_object.total = price40 + price200 + price1000 + \
                price_desalt + price_cartridge + price_setup
            row_object.total_degen = price40_degen + \
                price200_degen + price1000_degen
            billing_context.append(row_object)

        return render(request, 'oligos/billing_output.html', context={'billing': billing_context})


class EnterODView(FormView):
    def get(self, request):
        form = EnterODForm
        return render(request, 'oligos/enter_od.html', context={'form': form})

    def post(self, request):
        form = EnterODForm(request.POST)
        if not form.is_valid():
            return render(request, 'oligos/enter_od.html', context={'form': form})

        form_data = form.cleaned_data
        sample_volume = form_data['sample_volume']
        dilution_factor = form_data['dilution_factor']
        readings = form_data['readings']
        oligos_to_update = []
        for line in readings:
            line.strip().replace("\r", "")
            try:
                groups = re.match(
                    '^(\d+)[\t;,\s]+([\d\.]+).*$', line).groups()
                data_to_update = {
                    'id': groups[0],
                    'od_reading': groups[1],
                }
                oligos_to_update.append(data_to_update)
            except:
                messages.warning(
                    self.request, f'Unable to save {line}. Check for typos and try again.')
        if len(oligos_to_update) == 0:
            return render(request, 'oligos/enter_od.html', context={'form': form})

        save_count = 0
        saved_oligos = []
        for oligo_to_update in oligos_to_update:
            id = oligo_to_update['id']
            OD_reading = oligo_to_update['od_reading']
            try:
                oligo = Oligo.objects.get(id=id)
            except:
                messages.warning(
                    self.request, f'Could not find an oligo with ID {id}. Check ID and try again.')
                continue
            oligo.OD_reading = OD_reading
            oligo.volume = sample_volume
            oligo.OD_reading_dilution_factor = dilution_factor
            oligo.save()
            saved_oligos.append(oligo)
            save_count += 1

        if save_count > 0:
            messages.success(
                self.request, f'Saved {save_count} OD readings successfully.')

        saved_oligos.sort(key=lambda x: x.id)

        return render(request, 'oligos/od_saved.html', context={'oligos': saved_oligos})


class RemoveOrderView(FormView):
    form_class = IdRangeForm
    template_name = 'oligos/remove_order.html'

    def form_valid(self, form):
        data = form.cleaned_data

        if not data['low'] and not data['high']:
            messages.info(self.request, 'No data submitted.')
            return render(self.request, 'oligos/remove_order.html', context={'form': form})

        if not data['low'] or not data['high']:
            order_id = data['low'] if data['low'] else data['high']
            oligos_to_remove = list(Oligo.objects.filter(order_id=order_id))
        else:
            oligos_to_remove = Oligo.objects.filter(order_id__gte=data['low']).filter(
                order_id__lte=data['high']).order_by('order_id')

        if not oligos_to_remove:
            messages.info(self.request, 'No orders within that range.')
            return render(self.request, 'oligos/remove_order.html', context={'form': form})

        order_ids = [oligo.order_id for oligo in oligos_to_remove]

        if self.request.POST.get('submit') == 'Confirm and Remove':
            order_ids = list(dict.fromkeys(order_ids))
            for order in order_ids:
                oligos_in_order = Oligo.objects.filter(order_id=order)
                for oligo in oligos_in_order:
                    oligo.delete()
            messages.success(self.request, 'Removed orders successfully')
            return HttpResponseRedirect(reverse('oligos:client_order_list'))

        order_context = {}
        for order in order_ids:
            try:
                order_context[order] += 1
            except KeyError:
                order_context[order] = 1

        return render(self.request, 'oligos/remove_order.html', context={'form': form, 'orders': order_context, 'confirming': True})


def CreateLabels(oligos_id_list):
    path = 'static/template.docx'
    doc = Document(path)
    buffer = io.BytesIO()

    oligos = Oligo.objects.filter(id__in=oligos_id_list)

    table = doc.tables[0]

    label_text_list = []
    for oligo in oligos:
        label_text = f'{oligo.name} {oligo.created_at:%Y/%m/%d} GMS\n'
        label_text += f'{oligo.pmol_per_microliter} pmol/µL  {oligo.micrograms_per_microliter} µg/µL ID{oligo.id}\n'
        label_text += f'{oligo.sequence.lower()}'
        label_text_list.append(label_text)

    oligo_counter = 0
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i % 2 == 0:
                if oligo_counter < len(label_text_list):
                    cell.paragraphs[0].runs[0].text = label_text_list[oligo_counter]
                    oligo_counter += 1
                else:
                    cell.paragraphs[0].runs[0].text = ''

    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(buffer, as_attachment=True, filename="oligo_labels.docx")
